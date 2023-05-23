from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm, PasswordChangeForm
from django.contrib.auth import authenticate, login, update_session_auth_hash,logout
from django.contrib.auth.models import Group
from django.contrib import messages
from .models import ProUser, Donor, BloodRequest,Rejection,Invitation,Donation
from .forms import ProUserForm,LoginForm, DonorForm,BloodRequestForm,RejectionForm,InvitationForm,DonationForm
from django.core.mail import send_mail
from docx import Document
from io import BytesIO
from django.http import HttpResponse
from django.db.models import Sum,Avg,BooleanField, Case, When

def home(request):
    return render(request, 'registration/home.html')

def register(request):
    if request.method == 'POST':
        form = ProUserForm(request.POST)
        if form.is_valid():
            # Перевірка коду запрошення
            invitation_code = form.cleaned_data['invitation_code']
            try:
                invitation = Invitation.objects.get(token=invitation_code, is_active=True)
            except Invitation.DoesNotExist:
                # Якщо код запрошення недійсний або неактивний, перенаправлення на сторінку помилки або повідомлення
                return redirect('invitation_error')
            
            user = form.save()
            invitation.is_active = False
            invitation.save()
            login(request, user)
            return redirect('profile')
    else:
        form = ProUserForm()
    return render(request, 'registration/register.html', {'form': form})
def send_invitation_email(email, token):
    subject = 'Запрошення на сайт'
    message = f'Ви отримали запрошення на сайт. Ваш токен: {token}'
    from_email = 'anja.kuzmich@gmail.com'

    send_mail(subject, message, from_email, [email])

@login_required
def create_invitation(request):
    if request.method == 'POST':
        form = InvitationForm(request.POST, user=request.user)
        if form.is_valid():
            invitation = form.save(commit=False)
            invitation.sender = request.user
            invitation.save()

            # Відправка запрошення користувачу
            send_invitation_email(invitation.email, invitation.token)

            return redirect('invitation_success')
    else:
        form = InvitationForm(user=request.user)
    return render(request, 'registration/create_invitation.html', {'form': form})




def invitation_success(request):
    return render(request, 'registration/invitation_success.html')


@login_required
def profile(request):
    user = request.user
    return render(request, 'registration/profile.html', {'user': user})

@login_required
def edit_profile(request):
    user = request.user
    if request.method == 'POST':
        form = ProUserForm(request.POST, instance=user)
        if form.is_valid():
            new_username = form.cleaned_data['username']
            if new_username != user.username:
                form.add_error('username', 'You cannot change your username.')
                return render(request, 'registration/edit_profile.html', {'form': form})
            else:
                # Збереження форми без збереження юзернейму
                form.save(commit=False)
                form.instance.username = user.username
                form.save()
                return redirect('profile')
    else:
        form = ProUserForm(instance=user)
    return render(request, 'registration/edit_profile.html', {'form': form})

@login_required
def logout_user(request):
    logout(request)  # Викликайте функцію logout
    return redirect('home')

@login_required
def delete_profile(request):
    user = request.user
    user.delete()
    logout(request)
    messages.success(request, 'Your account was deleted successfully!')
    return redirect('registration/register')


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('profile')
    else:
        form = LoginForm()
    return render(request, 'registration/login.html', {'form': form})


def donor_list(request):
    # Отримати параметри сортування та пошуку з запиту
    sort_by = request.GET.get('sort_by', 'last_name')  # За замовчуванням сортувати за прізвищем
    search_query = request.GET.get('search_query', '')
    blood_group_filter = request.GET.get('blood_group', '')

    # Отримати список донорів з бази даних з врахуванням сортування та пошуку
    donors = Donor.objects.filter(first_name__icontains=search_query).order_by(sort_by)
    if blood_group_filter:
        donors = donors.filter(blood_group=blood_group_filter)

    context = {
        'donors': donors,
        'sort_by': sort_by,
        'search_query': search_query
    }
    return render(request, 'registration/donor_list.html', context)
    

def add_donor(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Перенаправте на сторінку входу, якщо користувач не аутентифікований

    form = DonorForm(request.POST or None,user=request.user)  # Передайте поточного користувача до форми

    if request.method == 'POST':
        if form.is_valid():
            form.save()
            return redirect('donor_list')  # Перенаправте на сторінку списку донорів після успішного створення

    return render(request, 'registration/add_donor.html', {'form': form})

@login_required
def create_blood_request(request, donor_id):
    donor = get_object_or_404(Donor, id=donor_id)

    if request.method == 'POST':
        form = BloodRequestForm(request.POST, initial={'donor': donor}, user=request.user)

        if form.is_valid():
            blood_request = form.save(commit=False)
            blood_request.user = request.user
            blood_request.save()
            if not blood_request.is_approved:
                return redirect('create_rejection', donation_id=form.instance.id)
            donor.donation_count = donor.donation_count + 1  # Додати 1 до кількості кроводачів донора
            donor.save()
            return redirect('donor_detail', donor_id)
    else:
        form = BloodRequestForm(initial={'donor': donor}, user=request.user)

    context = {
        'form': form,
        'donor_id': donor_id,
        'donation_id': form.instance.id,
    }
    return render(request, 'registration/blood_request_form.html', context)




def donor_detail(request, donor_id):
    donor = get_object_or_404(Donor, id=donor_id)
    return render(request, 'registration/donor_detail.html', {'donor': donor})


def donation_details(request, donation_id):
    donation = get_object_or_404(BloodRequest, id=donation_id)
    return render(request, 'registration/donation_details.html', {'donation': donation})


def donor_donations(request, donor_id):
    donor = Donor.objects.get(id=donor_id)
    donations = BloodRequest.objects.filter(donor=donor).order_by('-donation_date')

    # Пошук по даті
    search_query = request.GET.get('search', '')
    if search_query:
        donations = donations.filter(donation_date__icontains=search_query)

    return render(request, 'registration/donor_donations.html', {'donor': donor, 'donations': donations, 'search_query': search_query})

def create_rejection(request, donation_id):
    blood_request = get_object_or_404(BloodRequest, id=donation_id)
    
    if request.method == 'POST':
        form = RejectionForm(request.POST)
        if form.is_valid():
            reason = form.cleaned_data['reason']
            unavailability_term = form.cleaned_data['unavailability_term']
            
            # Створення запису про брак
            Rejection.objects.create(blood_request=blood_request, reason=reason, unavailability_term=unavailability_term)
            
            return redirect('show_rejections')  # Перенаправлення на список браків
    else:
        form = RejectionForm()
    
    context = {
        'form': form,
        'donation_id': donation_id
    }
    return render(request, 'registration/create_rejection.html', context)


def show_rejections(request):
    rejections = Rejection.objects.select_related('blood_request__donor')
    context = {'rejections': rejections}
    return render(request, 'registration/rejections.html', context)

# В'юшка для підтвердження донації
def confirm_donation(request, donation_id):
    donation = get_object_or_404(BloodRequest, id=donation_id)
    
    if request.method == 'POST':
        form = DonationForm(request.POST)
        
        if form.is_valid():
            # Отримайте дані з форми
            donation_type = form.cleaned_data['donation_type']
            plasma_units = form.cleaned_data['plasma_units']
            blood_units = form.cleaned_data['blood_units']
            
            # Створіть запис в таблиці Donation
            new_donation = Donation.objects.create(
                blood_request=donation,
                donation_date=donation.donation_date,
                donation_type=donation_type,
                plasma_units=plasma_units,
                blood_units=blood_units,
            )
            
            # Позначте початковий запит на кров як підтверджений
            donation.is_confirmed = True
            donation.save()
            
            # Перенаправте користувача на сторінку деталей донації або іншу відповідну сторінку
            return redirect('donation_details', donation_id=donation.id)
    else:
        form = DonationForm()
    print("form")
    return render(request, 'registration/confirm_donation.html', {'donation': donation, 'form': form})

def generate_report(request):
    donors = Donor.objects.all()
    # Кількість активних донорів
    permanent_donors = list(filter(lambda donor: donor.is_permanent_donor, donors))
    active_donors_count = len(permanent_donors)
    active_donors_donation_count = sum(donor.donation_count for donor in permanent_donors)

    print(active_donors_donation_count)
    reserve_donors = list(filter(lambda donor: donor.is_permanent_donor==False, donors))
    reserve_donors_count = len(reserve_donors)
    reserve_donors_donation_count = Donation.objects.filter(blood_request__donor__in=reserve_donors).count()
    plasma_donors_count = Donor.objects.filter(blood_requests__donation__donation_type='плазма').distinct().count()
    # Кількість імунних донорів
    immune_donors_count = Donor.objects.filter(blood_requests__donation__donation_type='імунна').distinct().count()

    # Кількість ізоімунних донорів
    isoimmune_donors_count = Donor.objects.filter(blood_requests__donation__donation_type='ізоімунна').distinct().count()

    # Кількість донорів клітин крові
    blood_cell_donors_count = Donor.objects.filter(blood_requests__donation__donation_type='донор клітин крові').distinct().count()

    # Загальна кількість донацій донорів (активних, резерву)
    total_donations_count = active_donors_donation_count + reserve_donors_donation_count

    # Загальна кількість кроводачів (без урахування плазмодачів)
    blood_donors_count = Donor.objects.filter(blood_requests__donation__donation_type__in=['імунна', 'ізоімунна', 'донор клітин крові', 'кров']).count()
    print(blood_donors_count)
    filtered_donors = list(filter(lambda donor: not donor.is_permanent_donor and BloodRequest.objects.filter(donor=donor, donation__donation_type__in=['кров']).exists(), donors))

    # Загальна кількість кроводачів від донорів резерву
    reserve_blood_donors_count =  Donation.objects.filter(blood_request__donor__in=filtered_donors).count()
    print(reserve_blood_donors_count)

    # Загальна кількість плазмодачів
    plasma_donors_count = Donor.objects.filter(blood_requests__donation__donation_type='плазма').count()
    print(plasma_donors_count)
    filtered_donors1 = list(filter(lambda donor: not donor.is_permanent_donor and BloodRequest.objects.filter(donor=donor, donation__donation_type__in=['плазма']).exists(), donors))
    reserve_plasma_donors_count = Donation.objects.filter(blood_request__donor__in=filtered_donors1).count()
    # Кількість плазмодач при однократному плазмаферезі
    plasma_donors_single_count = Donation.objects.filter(donation_type='плазма', plasma_units__gt=0).count()

    # Кількість плазмодач при двократному плазмаферезі
    plasma_donors_double_count = Donation.objects.filter(donation_type='плазма', plasma_units__gt=0).exclude(plasma_units=1).count()

    # Кількість плазмодач апаратним плазмаферезом
    plasma_donors_machine_count = Donation.objects.filter(donation_type='плазма', plasma_units__gt=0).filter(blood_request__is_full_donation=True).count()

    # Кількість імунних донорів
    

    reserve_blood_units = sum(donation.blood_units for donation in Donation.objects.filter(blood_request__donor__in=reserve_donors, blood_request__donation__donation_type='кров'))
    print(reserve_blood_units)
    # Отримання суми полів blood_units для записів від донорів резерву

    # Отримання загальної кількості крові та кількості записів донацій
    total_blood_units = Donation.objects.filter(donation_type='кров').aggregate(total=Sum('blood_units'))['total']
    average_donation = Donation.objects.filter(donation_type='кров').aggregate(average=Avg('blood_units'))['average']
    print(total_blood_units)
    print(average_donation)
     # Створення документа Word
    document = Document()
    
    # Додавання статистичних даних до документа
    document.add_heading('Статистика', level=1)
    
    table1 = document.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'

    # Заголовки стовпців
    table1.cell(0, 0).text = 'Найменування показника'
    table1.cell(0, 1).text = 'Значення'

    # Дані
    data1 = [
        ('Кількість активних донорів', active_donors_count),
        ('Кількість донацій активних донорів', active_donors_donation_count),
        ('Кількість донорів резерву', reserve_donors_count),
        ('Кількість донацій донорів резерву', reserve_donors_donation_count),
        ('Загальна кількість донорів (активних, резерву)', active_donors_count + reserve_donors_count),
        ('Кількість первинних донорів', active_donors_count),
        ('Кількість донорів плазми', plasma_donors_count),
        ('Кількість імунних донорів', immune_donors_count),
        ('Кількість ізоімунних донорів', isoimmune_donors_count),
        ('Кількість донорів клітин крові', blood_cell_donors_count),
        ('Загальна кількість донацій донорів (активних, резерву)', total_donations_count),
        ('Загальна кількість кроводачів (без урахування плазмодачів)', blood_donors_count),
        ('Загальна кількість кроводачів від донорів резерву', reserve_blood_donors_count),
        ('Загальна кількість плазмодач', plasma_donors_count),
        ('Загальна кількість плазмодач від донорів резерву', reserve_plasma_donors_count),
        ('Кількість плазмодач при однократному плазмаферезі', plasma_donors_single_count),
        ('Кількість плазмодач при двократному плазмаферезі', plasma_donors_double_count),
        ('Кількість плазмодач апаратним плазмаферезом', plasma_donors_machine_count)
    ]

    # Додавання даних до таблиці
    for i, (name, value) in enumerate(data1, start=1):
        row = table1.add_row().cells
        row[0].text = name
        row[1].text = str(value)
    
    
    # Додавання статистичних даних до документа
    document.add_paragraph()  # Додати порожній абзац між таблицями
    document.add_heading('Інша статистика', level=1)

    table2 = document.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'

    # Заголовки стовпців
    table2.cell(0, 0).text = 'Найменування показника'
    table2.cell(0, 1).text = 'Значення'

    # Дані
    data2 = [
        ('Заготовлено донорської крові',total_blood_units),
        ('Середня доза кроводачі, мл', average_donation),
        ('Заготовлено крові без гемоконсерванту від донорів резерву, л ', reserve_blood_units),]

    for i, (name, value) in enumerate(data2, start=1):
        row = table2.add_row().cells
        row[0].text = name
        row[1].text = str(value)

    # Збереження документа в буфер
    '''buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    # Відправлення документа як відповіді HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=statistics.docx'
    response.write(buffer.read())
    
    return response'''